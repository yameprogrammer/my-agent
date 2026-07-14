import io
import asyncio
from typing import List, Dict, Any
from sqlmodel import select
from sqlmodel.ext.asyncio.session import AsyncSession
from sqlalchemy.orm import selectinload

from app.models import Project, Episode, Content

# Import Guards
try:
    from ebooklib import epub
except ImportError:
    epub = None

try:
    from weasyprint import HTML
except ImportError:
    HTML = None

try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    docx = None

async def compile_novel_draft(project_id: int, db: AsyncSession) -> List[Dict[str, Any]]:
    """
    프로젝트 내의 에피소드를 episode_number 순으로 정렬하여,
    최종 승인된 본문(is_approved=True)을 추출합니다.
    승인본이 없을 경우, 가장 최근 생성된 버전을 Fallback으로 사용합니다.
    """
    stmt = (
        select(Episode)
        .where(Episode.project_id == project_id)
        .order_by(Episode.episode_number.asc())
        .options(selectinload(Episode.contents))
    )
    result = await db.execute(stmt)
    episodes = result.scalars().all()
    
    compiled_episodes = []
    for ep in episodes:
        # 1. 승인된 컨텐츠 탐색
        approved_content = next((c for c in ep.contents if c.is_approved), None)
        
        # 2. 없을 경우 최신 버전 Fallback 조회
        if not approved_content and ep.contents:
            sorted_contents = sorted(ep.contents, key=lambda c: c.created_at, reverse=True)
            approved_content = sorted_contents[0]
            
        text = approved_content.content_text if approved_content else "집필된 본문이 없습니다."
        compiled_episodes.append({
            "episode_number": ep.episode_number,
            "episode_title": ep.title,
            "text": text
        })
        
    return compiled_episodes

class NovelCompiler:
    def __init__(self, title: str, author: str, episodes: List[Dict[str, Any]]):
        self.title = title
        self.author = author
        self.episodes = episodes

    def build_txt(self) -> bytes:
        """
        소설 원고를 플랫폼 투고 형식 TXT로 조립합니다.
        """
        out = io.StringIO()
        out.write(f"제목: {self.title}\n")
        out.write(f"작가: {self.author}\n")
        out.write("=" * 40 + "\n\n")
        
        for ep in self.episodes:
            out.write(f"제 {ep['episode_number']}화. {ep['episode_title']}\n\n")
            out.write(ep["text"])
            out.write("\n\n\n◆ ◆ ◆\n\n\n")
            
        return out.getvalue().encode("utf-8")

    def build_epub(self) -> bytes:
        """
        소설 원고를 EPUB 규격으로 패킹합니다.
        """
        if not epub:
            raise ImportError("ebooklib 라이브러리가 설치되지 않았습니다. pip install ebooklib 실행 필요.")
            
        book = epub.EpubBook()
        book.set_identifier(f"novel-agentic-{self.title}")
        book.set_title(self.title)
        book.set_language("ko")
        book.add_author(self.author)
        
        # 기본 네비게이션 파일 등록
        spine_list = ['nav']
        toc_list = []
        
        for idx, ep in enumerate(self.episodes):
            chapter = epub.EpubHtml(
                title=ep["episode_title"],
                file_name=f"chap_{idx+1}.xhtml",
                lang="ko"
            )
            body_html = ep["text"].replace("\n", "<br/>")
            chapter.content = f"""
            <html>
            <head><title>{ep['episode_title']}</title></head>
            <body>
                <h2>제 {ep['episode_number']}화. {ep['episode_title']}</h2>
                <div class="novel-body" style="line-height: 1.8; font-size: 1.1em;">
                    {body_html}
                </div>
            </body>
            </html>
            """
            book.add_item(chapter)
            spine_list.append(chapter)
            toc_list.append(epub.Link(f"chap_{idx+1}.xhtml", ep["episode_title"], f"chap_{idx+1}"))
            
        book.spine = spine_list
        book.toc = tuple(toc_list)
        
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        out_stream = io.BytesIO()
        epub.write_epub(out_stream, book, {})
        return out_stream.getvalue()

    def build_pdf(self) -> bytes:
        """
        HTML 템플릿 + WeasyPrint 조합으로 PDF를 빌드합니다.
        """
        if not HTML:
            raise ImportError("weasyprint 라이브러리가 설치되지 않았습니다. pip install weasyprint 실행 필요.")
            
        episodes_html = ""
        for ep in self.episodes:
            body_html = ep["text"].replace("\n", "<br/>")
            episodes_html += f"""
            <div class="chapter">
                <h2 class="chapter-title">제 {ep['episode_number']}화. {ep['episode_title']}</h2>
                <div class="chapter-text">{body_html}</div>
            </div>
            """
            
        full_html = f"""
        <html>
        <head>
            <style>
                @page {{
                    size: A5;
                    margin: 25mm 20mm 25mm 20mm;
                    @bottom-center {{
                        content: counter(page);
                        font-family: serif;
                        font-size: 9pt;
                    }}
                }}
                body {{
                    font-family: 'KoPub Batang', 'Batang', serif;
                    line-height: 1.8;
                    font-size: 10.5pt;
                    text-align: justify;
                }}
                .title-page {{
                    page-break-after: always;
                    text-align: center;
                    margin-top: 50mm;
                }}
                .novel-title {{
                    font-size: 28pt;
                    font-weight: bold;
                    margin-bottom: 10mm;
                }}
                .novel-author {{
                    font-size: 16pt;
                    color: #555;
                }}
                .chapter {{
                    page-break-before: always;
                }}
                .chapter-title {{
                    font-size: 18pt;
                    text-align: center;
                    margin-bottom: 15mm;
                    margin-top: 10mm;
                }}
                .chapter-text {{
                    margin-top: 10mm;
                }}
            </style>
        </head>
        <body>
            <div class="title-page">
                <div class="novel-title">{self.title}</div>
                <div class="novel-author">저자: {self.author}</div>
            </div>
            {episodes_html}
        </body>
        </html>
        """
        out_stream = io.BytesIO()
        HTML(string=full_html).write_pdf(out_stream)
        return out_stream.getvalue()

    def build_docx(self) -> bytes:
        """
        python-docx 연동하여 Word 문서를 빌드합니다.
        """
        if not Document:
            raise ImportError("python-docx 라이브러리가 설치되지 않았습니다. pip install python-docx 실행 필요.")
            
        doc = Document()
        
        # 타이틀 페이지 구성
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(self.title)
        title_run.font.name = "Malgun Gothic"
        title_run.font.size = Pt(28)
        title_run.bold = True
        
        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_p.add_run(f"저자: {self.author}")
        author_run.font.size = Pt(14)
        author_run.font.color.rgb = docx.shared.RGBColor(100, 100, 100)
        
        doc.add_page_break()
        
        # 챕터 쓰기
        for ep in self.episodes:
            heading = doc.add_heading(f"제 {ep['episode_number']}화. {ep['episode_title']}", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.size = Pt(18)
            heading.runs[0].font.name = "Malgun Gothic"
            
            # 본문 작성
            for line in ep["text"].split("\n"):
                if not line.strip():
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Pt(10)
                p_run = p.add_run(line)
                p_run.font.size = Pt(11)
                p_run.font.name = "Batang"
                
            doc.add_page_break()
            
        out_stream = io.BytesIO()
        doc.save(out_stream)
        return out_stream.getvalue()
